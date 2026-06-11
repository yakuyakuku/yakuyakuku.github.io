import os
import sys
import re
import docx
from docx.oxml import parse_xml
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX

class MathParser:
    def parse(self, s):
        s = s.strip()
        if not s:
            return ""
            
        # Parse equality or basic operators (+, -, =, times ' x ') at paren depth 0
        depth = 0
        split_idx = -1
        split_op = None
        for i, char in enumerate(s):
            if char == '(':
                depth += 1
            elif char == ')':
                depth -= 1
            elif depth == 0:
                if char == '=':
                    split_idx = i
                    split_op = '='
                    break
                elif char == '+' and split_op not in ['=']:
                    split_idx = i
                    split_op = '+'
                elif char == '-' and split_op not in ['=', '+']:
                    split_idx = i
                    split_op = '-'
                elif s[i:i+3] == ' x ' and split_op not in ['=', '+', '-']:
                    split_idx = i
                    split_op = ' x '
                    
        if split_idx != -1:
            left = s[:split_idx]
            right = s[split_idx + len(split_op):]
            if split_op == ' x ':
                return self.parse(left) + "<m:r><m:t> \u00d7 </m:t></m:r>" + self.parse(right)
            elif split_op == '=':
                return self.parse(left) + "<m:r><m:t> = </m:t></m:r>" + self.parse(right)
            elif split_op == '+':
                return self.parse(left) + "<m:r><m:t> + </m:t></m:r>" + self.parse(right)
            elif split_op == '-':
                return self.parse(left) + "<m:r><m:t> - </m:t></m:r>" + self.parse(right)

        # Parse division '/' at paren depth 0
        depth = 0
        for i, char in enumerate(s):
            if char == '(':
                depth += 1
            elif char == ')':
                depth -= 1
            elif depth == 0 and char == '/':
                left = s[:i]
                right = s[i+1:]
                if left.startswith('(') and left.endswith(')'): left = left[1:-1]
                if right.startswith('(') and right.endswith(')'): right = right[1:-1]
                return f"<m:f><m:num>{self.parse(left)}</m:num><m:den>{self.parse(right)}</m:den></m:f>"

        # Parse powers '^' at paren depth 0 (right-to-left)
        depth = 0
        for i in range(len(s) - 1, -1, -1):
            char = s[i]
            if char == ')':
                depth += 1
            elif char == '(':
                depth -= 1
            elif depth == 0 and char == '^':
                base = s[:i]
                exponent = s[i+1:]
                if base.startswith('(') and base.endswith(')'): base = base[1:-1]
                if exponent.startswith('(') and exponent.endswith(')'): exponent = exponent[1:-1]
                return f"<m:sSup><m:e>{self.parse(base)}</m:e><m:sup>{self.parse(exponent)}</m:sup></m:sSup>"

        # Parse square root 'sqrt(...)' (Must start with sqrt and end with paren)
        if s.startswith('sqrt(') and s.endswith(')'):
            radicand = s[5:-1]
            # Word Office Math requires radPr and degHide to completely hide the root index box
            return f"<m:rad><m:radPr><m:degHide m:val=\"1\"/></m:radPr><m:deg/><m:e>{self.parse(radicand)}</m:e></m:rad>"
            
        # Parse implicit multiplication for sqrt (e.g., 3sqrt(2) -> 3 * sqrt(2))
        idx_sqrt = s.find("sqrt(")
        if idx_sqrt > 0:
            return self.parse(s[:idx_sqrt]) + self.parse(s[idx_sqrt:])
            
        # Parse implicit multiplication for parentheses (e.g., 2(x+3) -> 2 * (x+3))
        idx_paren = s.find("(")
        if idx_paren > 0:
            return self.parse(s[:idx_paren]) + self.parse(s[idx_paren:])

        # Parse text inside parenthesis
        if s.startswith('(') and s.endswith(')'):
            inner = s[1:-1]
            return f"<m:r><m:t>(</m:t></m:r>{self.parse(inner)}<m:r><m:t>)</m:t></m:r>"

        # Fallback to plain run
        t = s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        return f"<m:r><m:t>{t}</m:t></m:r>"

def format_text_in_paragraph(paragraph, english_regex, math_regex, bold_regex):
    text = paragraph.text
    if not text.strip():
        return
        
    matches = []
    
    # 1. Match Bold concepts
    for match in re.finditer(bold_regex, text):
        matches.append((match.start(), match.end(), 'bold'))
        
    # 2. Match English terms
    for match in re.finditer(english_regex, text, re.IGNORECASE):
        start, end = match.start(), match.end()
        # Prevent overlapping with bold concepts
        if not any(m[0] <= start < m[1] or m[0] < end <= m[1] for m in matches):
            matches.append((start, end, 'english'))
        
    # 3. Match Math expressions
    for match in re.finditer(math_regex, text):
        start, end = match.start(), match.end()
        # Prevent overlapping with existing matches
        if not any(m[0] <= start < m[1] or m[0] < end <= m[1] for m in matches):
            matches.append((start, end, 'math'))
            
    # Sort matches
    matches.sort(key=lambda x: x[0])
    
    if not matches:
        return
        
    # Clear runs in the paragraph
    p_element = paragraph._element
    for run in list(paragraph.runs):
        p_element.remove(run._element)
        
    # Also remove any existing oMath elements to be safe
    for child in list(p_element):
        if child.tag.endswith('oMath'):
            p_element.remove(child)
            
    last_idx = 0
    parser = MathParser()
    for start, end, match_type in matches:
        # Normal text run before match
        if start > last_idx:
            paragraph.add_run(text[last_idx:start])
            
        match_text = text[start:end]
        if match_type == 'bold':
            run = paragraph.add_run(match_text)
            run.bold = True
        elif match_type == 'english':
            # English term styled as italic
            run = paragraph.add_run(match_text)
            run.italic = True
        elif match_type == 'math':
            # Math equation parsed into OMML XML
            try:
                omml_body = parser.parse(match_text)
                omml_xml = f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{omml_body}</m:oMath>'
                omml_element = parse_xml(omml_xml)
                p_element.append(omml_element)
            except Exception as e:
                # Fallback to italic run on error
                print(f"Error parsing math '{match_text}': {e}")
                run = paragraph.add_run(match_text)
                run.italic = True
                
        last_idx = end
        
    # Remaining text
    if last_idx < len(text):
        paragraph.add_run(text[last_idx:])

def main():
    if len(sys.argv) < 2:
        print("Usage: python apply_formatting.py <output_docx_path>")
        sys.exit(1)
        
    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        print(f"Error: File not found at {file_path}")
        sys.exit(1)
        
    print(f"Applying native Office Math equations and English term formatting to: {file_path}")
    doc = docx.Document(file_path)
    
    # Combined regex patterns for English formal jargon
    english_terms = [
        r"expanding brackets", r"collecting like terms", r"factorising",
        r"rational exponents", r"laws of indices", r"rationalising the denominator",
        r"scaffolding", r"exit ticket", r"pre-test", r"flowchart",
        r"computational thinking", r"deep learning", r"differentiation",
        r"double check", r"self-reflection", r"peer evaluation", r"color-coded",
        r"data collection", r"data processing", r"verification", r"generalization",
        r"discovery learning", r"inquiry learning", r"problem based learning"
    ]
    english_regex = r"\b(?:" + "|".join(english_terms) + r")\b"

    # Bold concepts (case-sensitive or insensitive as required)
    bold_concepts = [
        r"Sains \(Fisika\)",
        r"Teknologi Informasi dan Rekayasa",
        r"Ekonomi"
    ]
    bold_regex = r"\b(?:" + "|".join(bold_concepts) + r")\b"
    
    # Math expressions: strict, non-greedy patterns. 
    # NOTE: Patterns ending in closing parenthesis ')' must NOT end in '\b' because ')' is a non-word character.
    math_patterns = [
        # Explicit equations/expressions with spaces strictly bound
        r"\b\d+\([x-z]\+\d+\)\s*-\s*\([x-z]-\d+\)",
        r"\b[x-z]\^2\s*\+\s*[a-z]*[x-z]\s*\+\s*[a-z]\b",
        r"\b[x-z]\^2\s*\+\s*\([a-z]\+[a-z]\)[x-z]\s*\+\s*[a-z]{2}\b",
        r"\b\d+\^\d+\s*x\s*\d+\^\d+\b",
        r"\b[a-z]\^\d+\s*=\s*\d+\b",
        r"\b[a-z]\^\(-[a-z]\)\s*=\s*\d+/\([a-z]\^[a-z]\)",
        r"\b[a-z]\^\(-?\d+/\d+\)\s*=\s*\d+/sqrt\([a-z]\)",
        r"\bsqrt\([a-z]\)\s*x\s*sqrt\([a-z]\)\s*=\s*[a-z]\b",
        r"\bsqrt\(\d+\)\s*=\s*\d*sqrt\(\d+\)",
        r"\b[a-z]/\([a-z]\s*\+\s*sqrt\([a-z]\)\)",
        r"\b[a-z]/\(sqrt\([a-z]\)\s*-\s*sqrt\([a-z]\)\)",
        r"\b[a-z]\^2\s*-\s*[a-z]\^2\b",
        r"\bsqrt\(\d+\)/sqrt\(\d+\)",
        r"\b\d+/sqrt\(\d+\)",
        r"\b\d+/\(\d+-sqrt\(\d+\)\)",
        r"\b\d+[a-z]\s*[x/*]\s*\d+[a-z]\b",
        
        # Binomial multiplication
        r"\([a-zA-Z\d]+[+-][a-zA-Z\d]+\)\([a-zA-Z\d]+[+-][a-zA-Z\d]+\)",
        
        # Exponential & Surds terms (without spaces)
        r"\b[a-zA-Z\d]+\^\(?[a-zA-Z\d_+-/]+\)?",
        r"\bsqrt\([a-zA-Z\d_+-/]+\)",
        
        # Single variables (x, y, z, a, b, c, n, m) standing alone as math variables
        r"\b[x-zabcdenm]\b"
    ]
    math_regex = r"(?:" + "|".join(math_patterns) + r")"
    
    # Process all paragraphs
    for para in doc.paragraphs:
        format_text_in_paragraph(para, english_regex, math_regex, bold_regex)
        
    # Process all tables (cells and cell paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    format_text_in_paragraph(para, english_regex, math_regex, bold_regex)

    # Format Pertemuan headings: yellow highlight, bold, underline, and add spacing before
    for para in doc.paragraphs:
        if re.match(r'^Pertemuan\s+\d+', para.text.strip()):
            # Add space before the paragraph to separate from previous content
            para.paragraph_format.space_before = Pt(18)
            # Apply bold, yellow highlight, and underline to all runs
            for run in para.runs:
                run.bold = True
                run.underline = True
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.save(file_path)
    print("Formatting applied successfully!")

if __name__ == "__main__":
    main()

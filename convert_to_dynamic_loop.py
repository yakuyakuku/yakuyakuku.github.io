import docx

def run_conversion():
    # 1. Load clean copy
    doc = docx.Document('template.docx')

    # 2. Truncate everything from "Pertemuan 2" onwards
    body = doc.element.body
    elements = list(body)
    found_p2 = False
    for element in elements:
        if not found_p2:
            if element.tag.endswith('p'):
                p = docx.text.paragraph.Paragraph(element, doc)
                if "Pertemuan 2" in p.text:
                    found_p2 = True
        if found_p2:
            body.remove(element)

    # 3. Add the outer loop opening tag at the very beginning (preserve styling by inserting before first para)
    doc.paragraphs[0].insert_paragraph_before("{FOR p IN pertemuan_list}")

    # 4. Modify the meeting header paragraph (which is now doc.paragraphs[1])
    # Keep the original run formatting (bold, color) but replace text
    for para in doc.paragraphs:
        if "Pertemuan 1" in para.text:
            # Clear all runs except the first, and set the first run's text
            if len(para.runs) > 0:
                para.runs[0].text = "Pertemuan {$p.pertemuan_number} – {$p.model_pembelajaran}"
                for r in para.runs[1:]:
                    r.text = ""
            else:
                para.text = "Pertemuan {$p.pertemuan_number} – {$p.model_pembelajaran}"
            break

    # 5. Table 0 replacement (TP) - preserve fonts, keep clean
    t0 = doc.tables[0]
    cell_tp = t0.rows[0].cells[1]
    cell_tp.text = ""
    run_tp = cell_tp.paragraphs[0].add_run("{FOR tp IN $p.tujuan_pembelajaran}{$tp}{END-FOR tp}")

    # 6. Table 1 replacement (Activities)
    t1 = doc.tables[1]
    mapping = {
        2: (0, 'deep_learning'),
        4: (1, 'differentiation'),
        5: (2, 'deep_learning'),
        6: (3, 'deep_learning'),
        8: (4, 'deep_learning'),
        9: (5, 'differentiation')
    }

    for row_idx, (act_idx, innov_field) in mapping.items():
        row = t1.rows[row_idx]
        
        # Waktu
        row.cells[0].text = ""
        row.cells[0].paragraphs[0].add_run(f"{{$p.aktivitas_pembelajaran[{act_idx}].waktu_menit}}'")
        
        # Sintaks
        row.cells[1].text = ""
        row.cells[1].paragraphs[0].add_run(f"{{$p.aktivitas_pembelajaran[{act_idx}].tahapan_sintaks}}")
        
        # Prosedur Guru (Cell 2 and 3)
        row.cells[2].text = ""
        row.cells[2].paragraphs[0].add_run(f"{{$p.aktivitas_pembelajaran[{act_idx}].prosedur_guru}}")
        row.cells[3].text = ""
        row.cells[3].paragraphs[0].add_run(f"{{$p.aktivitas_pembelajaran[{act_idx}].prosedur_guru}}")
        
        # Prosedur Siswa
        row.cells[4].text = ""
        row.cells[4].paragraphs[0].add_run(f"{{$p.aktivitas_pembelajaran[{act_idx}].prosedur_siswa}}")
        
        # Pertanyaan Kunci
        row.cells[5].text = ""
        row.cells[5].paragraphs[0].add_run(f"{{FOR pk IN $p.aktivitas_pembelajaran[{act_idx}].pertanyaan_kunci}}{{$pk}}\n{{END-FOR pk}}")
        
        # Muatan Inovatif (applying formatting logic)
        cell_innov = row.cells[6]
        cell_innov.text = ""
        p_innov = cell_innov.paragraphs[0]
        
        if innov_field == 'deep_learning':
            # Bold label run
            r_lbl = p_innov.add_run(f"{{$p.aktivitas_pembelajaran[{act_idx}].muatan_inovatif.deep_learning ? 'Deep Learning: ' : ''}}")
            r_lbl.bold = True
            # Plain description run
            r_desc = p_innov.add_run(f"{{$p.aktivitas_pembelajaran[{act_idx}].muatan_inovatif.deep_learning || ''}}")
            
        elif innov_field == 'differentiation':
            # Bold label run
            r_lbl = p_innov.add_run(f"{{$p.aktivitas_pembelajaran[{act_idx}].muatan_inovatif.differentiation ? 'Differentiation (' + $p.aktivitas_pembelajaran[{act_idx}].muatan_inovatif.differentiation.type + '): ' : ''}}")
            r_lbl.bold = True
            # Plain description run
            r_desc = p_innov.add_run(f"{{$p.aktivitas_pembelajaran[{act_idx}].muatan_inovatif.differentiation ? $p.aktivitas_pembelajaran[{act_idx}].muatan_inovatif.differentiation.deskripsi : ''}}")
            
        else:
            p_innov.text = ""

    # 7. Table 2 replacement (Asesmen)
    t2 = doc.tables[2]
    for i in range(4):
        row = t2.rows[i + 1]
        
        # Cell 0 (Jenis Asesmen) - Bold
        cell_jenis = row.cells[0]
        cell_jenis.text = ""
        r_jenis = cell_jenis.paragraphs[0].add_run(f"{{$p.asesmen[{i}].jenis}}")
        r_jenis.bold = True
        
        # Cell 1 (Bentuk Asesmen) - Plain
        row.cells[1].text = f"{{$p.asesmen[{i}].bentuk}}"
        
        # Cell 2 (Lampiran Asesmen) - Bold label + Plain description
        cell_lampiran = row.cells[2]
        cell_lampiran.text = ""
        p_lampiran = cell_lampiran.paragraphs[0]
        
        # Bold label run
        r_lbl = p_lampiran.add_run(f"{{$p.asesmen[{i}].lampiran.includes(':') ? $p.asesmen[{i}].lampiran.split(':')[0] + ':' : ''}}")
        r_lbl.bold = True
        # Plain description run
        r_desc = p_lampiran.add_run(f"{{$p.asesmen[{i}].lampiran.includes(':') ? $p.asesmen[{i}].lampiran.substring($p.asesmen[{i}].lampiran.indexOf(':') + 1) : $p.asesmen[{i}].lampiran}}")

    # 8. Add the outer loop closing tag at the very end
    doc.add_paragraph("{END-FOR p}")

    doc.save('template.docx')
    print("Successfully converted template.docx with bold prefixes and conditional formatting!")

if __name__ == "__main__":
    run_conversion()

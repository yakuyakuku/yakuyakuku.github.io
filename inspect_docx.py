import docx

doc = docx.Document('template.docx')

print("--- Paragraphs ---")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"P{i}: {para.text[:100]}")

print("\n--- Tables ---")
for i, table in enumerate(doc.tables):
    print(f"Table {i}: {len(table.rows)} rows, {len(table.columns)} columns")
    # Print first cell of header and first cell of data row to identify it
    if len(table.rows) > 0:
        header = [cell.text.strip() for cell in table.rows[0].cells]
        print(f"  Header: {header}")
    if len(table.rows) > 1:
        first_data = [cell.text.strip()[:50] for cell in table.rows[1].cells]
        print(f"  Row 1: {first_data}")

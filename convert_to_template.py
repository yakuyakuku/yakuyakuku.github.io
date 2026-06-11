import docx

def run_conversion():
    doc = docx.Document('template.docx')

    # Paragraph replacement
    for para in doc.paragraphs:
        if "Pertemuan 1" in para.text:
            para.text = "Pertemuan {pertemuan_number} – {model_pembelajaran}"

    # Table 0 replacement (TP)
    t0 = doc.tables[0]
    t0.rows[0].cells[1].text = "{FOR tp IN tujuan_pembelajaran}{$tp}{END-FOR tp}"

    # Table 1 replacement (Activities)
    t1 = doc.tables[1]
    mapping = {
        2: (0, 'deep_learning'),
        4: (1, 'differentiation.deskripsi'),
        5: (2, None),
        6: (3, None),
        8: (4, None),
        9: (5, 'differentiation.deskripsi')
    }

    for row_idx, (act_idx, innov_field) in mapping.items():
        row = t1.rows[row_idx]
        row.cells[0].text = f"{{aktivitas_pembelajaran[{act_idx}].waktu_menit}}'"
        row.cells[1].text = f"{{aktivitas_pembelajaran[{act_idx}].tahapan_sintaks}}"
        
        # Prosedur Guru
        row.cells[2].text = f"{{aktivitas_pembelajaran[{act_idx}].prosedur_guru}}"
        row.cells[3].text = f"{{aktivitas_pembelajaran[{act_idx}].prosedur_guru}}"
        
        # Prosedur Siswa
        row.cells[4].text = f"{{aktivitas_pembelajaran[{act_idx}].prosedur_siswa}}"
        
        # Pertanyaan Kunci (using loop with $pk)
        row.cells[5].text = f"{{FOR pk IN aktivitas_pembelajaran[{act_idx}].pertanyaan_kunci}}{{$pk}}\n{{END-FOR pk}}"
        
        # Muatan Inovatif
        if innov_field == 'deep_learning':
            row.cells[6].text = f"{{aktivitas_pembelajaran[{act_idx}].muatan_inovatif.deep_learning}}"
        elif innov_field == 'differentiation.deskripsi':
            # Use safe JS conditional expression to avoid TypeError if differentiation is null
            row.cells[6].text = f"{{aktivitas_pembelajaran[{act_idx}].muatan_inovatif.differentiation ? 'Differentiation (' + aktivitas_pembelajaran[{act_idx}].muatan_inovatif.differentiation.type + '): ' + aktivitas_pembelajaran[{act_idx}].muatan_inovatif.differentiation.deskripsi : ''}}"
        else:
            row.cells[6].text = ""

    # Table 2 replacement (Asesmen)
    t2 = doc.tables[2]
    for i in range(4):
        row = t2.rows[i + 1]
        row.cells[0].text = f"{{asesmen[{i}].jenis}}"
        row.cells[1].text = f"{{asesmen[{i}].bentuk}}"
        row.cells[2].text = f"{{asesmen[{i}].lampiran}}"

    doc.save('template.docx')
    print("Successfully converted template.docx with safe conditional differentiation expressions!")

if __name__ == "__main__":
    run_conversion()

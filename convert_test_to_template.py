import docx

def run_conversion():
    # 1. Load clean copy
    doc = docx.Document('Test.docx')
    body = doc.element.body

    # 2. Truncate everything from "Pertemuan 2" up to "Pontianak"
    elements = list(body)
    found_p2 = False
    nodes_to_delete = []
    for element in elements:
        tag = element.tag.split('}')[-1]
        if not found_p2:
            if tag == 'p':
                p = docx.text.paragraph.Paragraph(element, doc)
                if p.text.strip().startswith("Pertemuan 2"):
                    found_p2 = True
                    nodes_to_delete.append(element)
        else:
            if tag == 'p':
                p = docx.text.paragraph.Paragraph(element, doc)
                if p.text.strip().startswith("Pontianak"):
                    break
            nodes_to_delete.append(element)

    for node in nodes_to_delete:
        body.remove(node)

    # 3. Add the outer loop opening tag before Pertemuan 1 and make it bold
    for p in doc.paragraphs:
        if p.text.strip().startswith("Pertemuan 1"):
            p.insert_paragraph_before("{FOR p IN pertemuan_list}")
            # Clear text and add bold run
            p.text = ""
            r = p.add_run("Pertemuan {$p.pertemuan_number} – {$p.model_pembelajaran}")
            r.bold = True
            break

    # 4. Add the outer loop closing tag before Pontianak and replace with dynamic date variable
    for p in doc.paragraphs:
        if p.text.strip().startswith("Pontianak"):
            p.insert_paragraph_before("{END-FOR p}")
            p.text = "Pontianak, {tanggal_generasi}"
            break

    # 5. Table 0 replacement (General Info)
    t0 = doc.tables[0]
    t0.rows[4].cells[1].text = "{judul_tema_subtopik}"

    # 6. Table 1 replacement (Konteks Pembelajaran)
    t1 = doc.tables[1]
    
    # Latar Belakang (Row 2, Cell 1)
    cell_lb = t1.rows[2].cells[1]
    cell_lb.text = "Latar belakang pembelajaran:\n{FOR lb IN latar_belakang_pembelajaran}• {$lb}\n{END-FOR lb}"
    
    # Capaian Pembelajaran (Row 5, Cell 1)
    cell_cp = t1.rows[5].cells[1]
    cell_cp.text = "{capaian_pembelajaran}"
    
    # Tujuan Pembelajaran Umum (Row 6, Cell 1)
    cell_tpu = t1.rows[6].cells[1]
    cell_tpu.text = "{FOR tpu IN tujuan_pembelajaran_umum}{$tpu}\n{END-FOR tpu}"
    
    # Enduring Understanding (Row 7, Cell 1)
    cell_eu = t1.rows[7].cells[1]
    cell_eu.text = "{enduring_understanding}"
    
    # Lintas Disiplin Ilmu (Row 8, Cell 1)
    cell_ld = t1.rows[8].cells[1]
    cell_ld.text = "{lintas_disiplin_ilmu}"

    # 7. Table 2 replacement (Sasaran Profil)
    t2 = doc.tables[2]
    
    # Impactful Checkboxes (Row 1, Cell 1)
    cell_impact = t2.rows[1].cells[1]
    cell_impact.text = ""
    impactful_values = [
        ("integrity_checked", "Integrity"),
        ("mindful_checked", "Mindful"),
        ("progressive_checked", "Progressive"),
        ("agility_checked", "Agility"),
        ("compassion_checked", "Compassion"),
        ("tenacity_checked", "Tenacity"),
        ("fidelity_checked", "Fidelity"),
        ("uplifting_checked", "Uplifting"),
        ("lifelong_learner_checked", "Lifelong Learner.")
    ]
    for i, (key, label) in enumerate(impactful_values):
        p = cell_impact.add_paragraph() if i > 0 else cell_impact.paragraphs[0]
        p.text = f"{{{key} ? '☑' : '☐'}} {label}"
    
    # Indikator (Row 1, Cell 2)
    cell_ind = t2.rows[1].cells[2]
    cell_ind.text = ""
    
    p1 = cell_ind.paragraphs[0]
    p1.text = "{FOR val IN sasaran_profil_sekolah}"
    
    p2 = cell_ind.add_paragraph()
    r2 = p2.add_run("{$val.index}. {$val.nilai}:")
    r2.bold = True
    
    p3 = cell_ind.add_paragraph()
    p3.text = "{FOR ind IN $val.indikator}"
    
    p4 = cell_ind.add_paragraph()
    p4.text = "• {$ind}"
    
    p5 = cell_ind.add_paragraph()
    p5.text = "{END-FOR ind}"
    
    p6 = cell_ind.add_paragraph()
    p6.text = "{END-FOR val}"

    # 8. Table 3 replacement (Pertemuan 1 TP)
    t3 = doc.tables[3]
    t3.rows[0].cells[1].text = "{FOR tp IN $p.tujuan_pembelajaran}{$tp}{END-FOR tp}"

    # 9. Table 4 replacement (Pertemuan 1 Activities)
    t4 = doc.tables[4]
    mapping = {
        2: 0,
        4: 1,
        5: 2,
        6: 3,
        8: 4,
        9: 5
    }

    for row_idx, act_idx in mapping.items():
        row = t4.rows[row_idx]
        
        # Waktu
        row.cells[0].text = f"{{$p.aktivitas_pembelajaran[{act_idx}].waktu_menit}}'"
        
        # Sintaks
        row.cells[1].text = f"{{$p.aktivitas_pembelajaran[{act_idx}].tahapan_sintaks}}"
        
        # Prosedur Guru (Cell 2 & 3)
        row.cells[2].text = f"{{$p.aktivitas_pembelajaran[{act_idx}].prosedur_guru}}"
        row.cells[3].text = f"{{$p.aktivitas_pembelajaran[{act_idx}].prosedur_guru}}"
        
        # Prosedur Siswa
        row.cells[4].text = f"{{$p.aktivitas_pembelajaran[{act_idx}].prosedur_siswa}}"
        
        # Pertanyaan Kunci
        row.cells[5].text = f"{{FOR pk IN $p.aktivitas_pembelajaran[{act_idx}].pertanyaan_kunci}}{{$pk}}\n{{END-FOR pk}}"
        
        # Muatan Inovatif - fully dynamic for all fields
        cell_innov = row.cells[6]
        cell_innov.text = ""
        p_innov = cell_innov.paragraphs[0]
        
        # Deep Learning run
        r_dl_lbl = p_innov.add_run(f"{{$p.aktivitas_pembelajaran[{act_idx}].muatan_inovatif.deep_learning ? 'Deep Learning: ' : ''}}")
        r_dl_lbl.bold = True
        p_innov.add_run(f"{{$p.aktivitas_pembelajaran[{act_idx}].muatan_inovatif.deep_learning || ''}}")
        
        # Differentiation run
        r_diff_lbl = p_innov.add_run(f"{{$p.aktivitas_pembelajaran[{act_idx}].muatan_inovatif.differentiation ? 'Differentiation (' + $p.aktivitas_pembelajaran[{act_idx}].muatan_inovatif.differentiation.type + '): ' : ''}}")
        r_diff_lbl.bold = True
        p_innov.add_run(f"{{$p.aktivitas_pembelajaran[{act_idx}].muatan_inovatif.differentiation ? $p.aktivitas_pembelajaran[{act_idx}].muatan_inovatif.differentiation.deskripsi : ''}}")
        
        # Impactful run
        r_imp_lbl = p_innov.add_run(f"{{$p.aktivitas_pembelajaran[{act_idx}].muatan_inovatif.impactful ? 'Impactful: ' : ''}}")
        r_imp_lbl.bold = True
        p_innov.add_run(f"{{$p.aktivitas_pembelajaran[{act_idx}].muatan_inovatif.impactful || ''}}")

    # 10. Table 5 replacement (Pertemuan 1 Asesmen)
    t5 = doc.tables[5]
    for i in range(4):
        row = t5.rows[i + 1]
        
        # Jenis
        cell_jenis = row.cells[0]
        cell_jenis.text = ""
        r_jenis = cell_jenis.paragraphs[0].add_run(f"{{$p.asesmen[{i}].jenis}}")
        r_jenis.bold = True
        
        # Bentuk
        row.cells[1].text = f"{{$p.asesmen[{i}].bentuk}}"
        
        # Lampiran (Bold label + plain description)
        cell_lamp = row.cells[2]
        cell_lamp.text = ""
        p_lamp = cell_lamp.paragraphs[0]
        
        r_lbl = p_lamp.add_run(f"{{$p.asesmen[{i}].lampiran.includes(':') ? $p.asesmen[{i}].lampiran.split(':')[0] + ':' : ''}}")
        r_lbl.bold = True
        r_desc = p_lamp.add_run(f"{{$p.asesmen[{i}].lampiran.includes(':') ? $p.asesmen[{i}].lampiran.substring($p.asesmen[{i}].lampiran.indexOf(':') + 1) : $p.asesmen[{i}].lampiran}}")

    # 11. Delete all elements after the LAMPIRAN paragraph, leaving only the header and a page break
    lampiran_element = None
    for child in list(body):
        if child.tag.split('}')[-1] == 'p':
            p = docx.text.paragraph.Paragraph(child, doc)
            if p.text.strip() == "LAMPIRAN":
                lampiran_element = child
                break
                
    if lampiran_element is not None:
        siblings = list(body)
        idx_lampiran = siblings.index(lampiran_element)
        for element in siblings[idx_lampiran + 1:]:
            body.remove(element)
        # Add a page break to create an empty page below LAMPIRAN
        doc.add_page_break()

    # 12. Delete empty spacing paragraphs inside the meeting loop to keep the layout tight
    loop_start_el = None
    loop_end_el = None
    for child in list(body):
        if child.tag.split('}')[-1] == 'p':
            p = docx.text.paragraph.Paragraph(child, doc)
            if p.text.strip() == "{FOR p IN pertemuan_list}":
                loop_start_el = child
            elif p.text.strip() == "{END-FOR p}":
                loop_end_el = child
                
    if loop_start_el is not None and loop_end_el is not None:
        siblings = list(body)
        idx_start = siblings.index(loop_start_el)
        idx_end = siblings.index(loop_end_el)
        for element in siblings[idx_start + 1 : idx_end]:
            if element.tag.split('}')[-1] == 'p':
                p = docx.text.paragraph.Paragraph(element, doc)
                if p.text.strip() == "":
                    body.remove(element)

    doc.save('template.docx')
    print("Successfully generated template.docx from Test.docx with comprehensive dynamic placeholders!")

if __name__ == "__main__":
    run_conversion()
